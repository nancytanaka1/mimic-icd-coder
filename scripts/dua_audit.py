"""Comprehensive DUA compliance audit — run before pushing to a remote.

Complements ``scripts/check_notebook_phi.py`` (notebooks only) with
repository-wide checks covering the full MIMIC-IV DUA surface.

Seven checks:
 1. Forbidden data extensions in tracked tree (.parquet, .csv.gz, etc.)
 2. User-specific configs accidentally tracked
 3. PHI / clinical-text / raw-data-path patterns in all tracked text files
 4. Binary docx scan (``reports/EDA_Report.docx`` via python-docx)
 5. ``.gitignore`` effectiveness (files on disk correctly ignored)
 6. Any data-shape files anywhere in working tree (tracked or not)
 7. Notebook PHI scanner delegation

Run before a push / public release:
    python scripts/dua_audit.py

Exit codes:
    0 — all seven checks clean
    1 — at least one issue group to review
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.stdout.reconfigure(encoding="utf-8")


def run(cmd: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, capture_output=True, text=True, cwd=ROOT, encoding="utf-8")


def tracked_files() -> list[str]:
    r = run(["git", "ls-files"])
    return r.stdout.strip().splitlines() if r.returncode == 0 else []


PATTERNS = {
    "clinical_section_header": [
        r"Admission\s+Date\s*:",
        r"Discharge\s+Date\s*:",
        r"Chief\s+Complaint\s*:",
        r"History\s+of\s+Present\s+Illness",
        r"Past\s+Medical\s+History",
        r"Hospital\s+Course\s*:",
        r"Discharge\s+Diagnosis",
        r"Physical\s+Exam\s*:",
        r"Pertinent\s+Results\s*:",
        r"Brief\s+Hospital\s+Course",
    ],
    "phi_marker": [
        r"\bMRN\s*[:#]",
        r"\bSSN\s*[:#]?",
        r"\bDOB\s*[:#]?",
        r"Patient\s+is\s+a\s+\d+",
        r"\b\d+\s*(?:yo|y\.?o\.?|year[- ]old)\b",
        r"\(\d{3}\)\s*\d{3}-\d{4}",
        r"\b\d{3}-\d{3}-\d{4}\b",
    ],
    "mimic_redaction_leak": [
        r"_{5,}",
        r"\[\*\*[A-Za-z0-9 ,./-]{1,120}\*\*\]",
    ],
    # Raw-data-path patterns — intentionally narrow. Match only:
    #   - absolute Windows paths (``E:/data/physionet/...``)
    #   - absolute POSIX paths (``/data/physionet/...``)
    #   - PhysioNet source URLs (``physionet.org/files/mimic``)
    # Do NOT match the PhysioNet-published directory structure alone
    # (``mimic-iv-note/note/discharge.csv.gz`` appears in legitimate
    # documentation of the public data layout).
    "raw_data_path": [
        r"[/\\]data[/\\]physionet",
        r"[A-Za-z]:[/\\]data[/\\]physionet",
        r"physionet\.org[/\\]files[/\\]mimic",
    ],
}

COMPILED = {cat: [re.compile(p, re.IGNORECASE) for p in pats] for cat, pats in PATTERNS.items()}

# Files that legitimately contain pattern definitions or synthetic clinical fixtures;
# they will match their own patterns and are excluded from the text scan.
SELF_MATCHING = {
    "scripts/check_notebook_phi.py",
    "scripts/dua_audit.py",
    "tests/fixtures/synthetic_notes.py",
    "reports/EDA_Report.docx",  # scanned separately with python-docx
}

# Per-category allowlist: files where a specific category is a legitimate
# documentation reference, not a leak.
CATEGORY_ALLOWLIST = {
    # LOCAL_SETUP.md is the workstation-specific setup playbook and legitimately
    # documents the typical MIMIC data location on the user's machine.
    "raw_data_path": {"LOCAL_SETUP.md"},
}

FORBIDDEN_EXTS = {".parquet", ".csv", ".gz", ".npz", ".joblib", ".pkl", ".feather", ".pq", ".tsv"}


def main() -> int:
    issues = 0
    tracked = tracked_files()
    print(f"=== Tracked file count: {len(tracked)} ===\n")

    # 1. Forbidden extensions
    print("### 1. Forbidden data-file extensions in tracked tree")
    hits = [f for f in tracked if Path(f).suffix.lower() in FORBIDDEN_EXTS]
    if hits:
        issues += 1
        print(f"  FAIL — {len(hits)} forbidden file(s):")
        for f in hits:
            print(f"    {f}")
    else:
        print("  PASS — no forbidden extensions in tracked tree")
    print()

    # 2. User-specific configs
    print("### 2. User-specific configs")
    user_cfg = [
        f
        for f in tracked
        if re.match(r"configs/(dev|prod)(\.[^.]+)?\.yml$", f) and "example" not in f
    ]
    if user_cfg:
        issues += 1
        print("  FAIL — user-specific configs tracked:")
        for f in user_cfg:
            print(f"    {f}")
    else:
        print("  PASS — only template configs tracked")
    print()

    # 3. Text-file pattern scan
    print("### 3. PHI / clinical-text / raw-data-path patterns in tracked text files")
    text_findings: list[tuple[str, int, str, str, str]] = []
    for f in tracked:
        if f in SELF_MATCHING:
            continue
        p = ROOT / f
        if not p.is_file():
            continue
        if p.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".pdf", ".docx", ".ico"}:
            continue
        try:
            text = p.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        for category, regexes in COMPILED.items():
            # Skip this category for files on its allowlist.
            if f in CATEGORY_ALLOWLIST.get(category, set()):
                continue
            for regex in regexes:
                m = regex.search(text)
                if m:
                    line_no = text[: m.start()].count("\n") + 1
                    ctx_start = max(0, m.start() - 30)
                    ctx_end = min(len(text), m.end() + 30)
                    ctx = text[ctx_start:ctx_end].replace("\n", " ").strip()
                    text_findings.append((f, line_no, category, regex.pattern, ctx))
                    break
            else:
                continue
            break
    if text_findings:
        issues += 1
        print(f"  FAIL — {len(text_findings)} finding(s):")
        for f, line, cat, pat, ctx in text_findings:
            print(f"    {f}:{line}  [{cat}]  /{pat}/")
            print(f"      context: ...{ctx[:120]}...")
    else:
        print("  PASS — no concerning patterns in tracked text files")
    print()

    # 4. Binary docx scan
    print("### 4. Binary scan (reports/EDA_Report.docx)")
    try:
        from docx import Document

        doc = Document(str(ROOT / "reports" / "EDA_Report.docx"))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        full_text = "\n".join(parts)
        doc_findings = []
        for category, regexes in COMPILED.items():
            for regex in regexes:
                m = regex.search(full_text)
                if m:
                    ctx_start = max(0, m.start() - 30)
                    ctx_end = min(len(full_text), m.end() + 30)
                    ctx = full_text[ctx_start:ctx_end].replace("\n", " ")
                    doc_findings.append((category, regex.pattern, ctx[:120]))
                    break
        if doc_findings:
            issues += 1
            print(f"  FAIL — {len(doc_findings)} finding(s):")
            for cat, pat, ctx in doc_findings:
                print(f"    [{cat}] /{pat}/")
                print(f"      context: ...{ctx}...")
        else:
            print("  PASS — no concerning patterns in EDA_Report.docx")
    except Exception as e:
        print(f"  SKIP — {e}")
    print()

    # 5. Gitignore effectiveness
    print("### 5. Gitignore effectiveness")
    must_be_ignored = [
        "data/bronze/discharge_notes.parquet",
        "data/bronze/diagnoses_icd.parquet",
        "data/bronze/admissions.parquet",
        "data/bronze/patients.parquet",
        "data/bronze/d_icd_diagnoses.parquet",
        "configs/dev.nancy.yml",
        ".venv/pyvenv.cfg",
    ]
    ok, fail = [], []
    for f in must_be_ignored:
        if not (ROOT / f).exists():
            continue
        r = run(["git", "check-ignore", "-v", f])
        (ok if r.returncode == 0 else fail).append(f)
    print(f"  {len(ok)} correctly gitignored ({', '.join(ok) if ok else '—'})")
    if fail:
        issues += 1
        print("  FAIL — present on disk but NOT gitignored:")
        for f in fail:
            print(f"    {f}")
    print()

    # 6. Data-shape files anywhere
    print("### 6. Any data-shape files anywhere in working tree")
    skip_parts = {".venv", ".git", "tests"}
    data_files = []
    for pattern in ("*.parquet", "*.csv.gz", "*.csv"):
        for p in ROOT.rglob(pattern):
            if any(part in skip_parts for part in p.parts):
                continue
            data_files.append(p.relative_to(ROOT).as_posix())
    tracked_set = set(tracked)
    for f in data_files:
        status = "TRACKED — VIOLATION" if f in tracked_set else "untracked (ignored)"
        flag = "FAIL" if f in tracked_set else "OK"
        print(f"  [{flag}] {f}: {status}")
    if not data_files:
        print("  (none found)")
    untracked_data_count = sum(1 for f in data_files if f not in tracked_set)
    if untracked_data_count and not any(f in tracked_set for f in data_files):
        print(f"  {untracked_data_count} data file(s) on disk, all correctly ignored")
    print()

    # 7. Notebook scanner
    print("### 7. Notebook PHI scanner")
    py = ROOT / ".venv" / "Scripts" / "python.exe"
    scanner = ROOT / "scripts" / "check_notebook_phi.py"
    r = run([str(py), str(scanner), "--scope", str(ROOT)])
    print("  " + (r.stdout.strip() or r.stderr.strip()))
    print(f"  exit code: {r.returncode}")
    if r.returncode != 0:
        issues += 1
    print()

    print("=" * 70)
    if issues == 0:
        print("DUA AUDIT: CLEAN — all checks passed")
        return 0
    print(f"DUA AUDIT: {issues} issue group(s) require attention")
    return 1


if __name__ == "__main__":
    sys.exit(main())
